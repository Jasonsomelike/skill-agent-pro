import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from PIL import Image

from skills.besti_document import generate_besti_docx


class BestiDocumentTests(unittest.TestCase):
    def test_generates_besti_docx_with_expected_geometry_and_fonts(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            output = root / "Besti测试公文.docx"
            image_path = root / "示意图.png"
            Image.new("RGB", (640, 240), "white").save(image_path)
            result = generate_besti_docx(
                root,
                output_path=output,
                title="关于开展计算机网络学习活动的通知",
                recipient="各学习小组",
                body=(
                    "为提升计算机网络课程学习效果，现将有关事项通知如下。\n"
                    "一、学习内容\n"
                    "（一）掌握 CIDR 与子网划分。\n"
                    "1. 完成三道练习题。"
                ),
                organization="院学习部",
                date_text="2026 年 6 月 18 日",
                attachments=["学习任务安排表"],
                images=[
                    {
                        "path": image_path.name,
                        "caption": "图 1  网络示意图",
                        "width_cm": 10,
                    }
                ],
                tables=[
                    {
                        "title": "学习任务安排表",
                        "headers": ["序号", "任务", "完成标准"],
                        "rows": [["1", "CIDR 练习", "独立完成三题"]],
                        "column_widths_cm": [2.0, 6.0, 7.6],
                    }
                ],
            )

            self.assertTrue(output.is_file())
            self.assertGreater(result["size"], 1000)
            document = Document(output)
            section = document.sections[0]
            self.assertAlmostEqual(section.top_margin.cm, 3.6, places=1)
            self.assertAlmostEqual(section.bottom_margin.cm, 3.0, places=1)
            self.assertAlmostEqual(section.left_margin.cm, 2.7, places=1)
            self.assertAlmostEqual(section.right_margin.cm, 2.7, places=1)

            title_run = next(
                paragraph.runs[0]
                for paragraph in document.paragraphs
                if paragraph.text == "关于开展计算机网络学习活动的通知"
            )
            self.assertEqual(title_run.font.size.pt, 22)
            self.assertEqual(
                title_run._element.rPr.rFonts.get(qn("w:eastAsia")),
                "方正小标宋简体",
            )
            self.assertIn("1．完成三道练习题。", [p.text for p in document.paragraphs])
            footer_xml = section.footer._element.xml
            self.assertIn(" PAGE ", footer_xml)
            body_children = list(document._element.body)
            signature_index = next(
                index
                for index, child in enumerate(body_children)
                if child.tag.endswith("}p") and "院学习部" in child.text
            )
            table_index = next(
                index
                for index, child in enumerate(body_children)
                if child.tag.endswith("}tbl")
            )
            self.assertLess(signature_index, table_index)
            attachment_paragraph = next(
                paragraph for paragraph in document.paragraphs if paragraph.text == "附件 1"
            )
            self.assertTrue(attachment_paragraph.paragraph_format.page_break_before)
            self.assertIn(
                "　　　2．图 1  网络示意图",
                [paragraph.text for paragraph in document.paragraphs],
            )
            image_paragraph = next(
                paragraph
                for paragraph in document.paragraphs
                if "<pic:pic" in paragraph._element.xml
            )
            self.assertNotIn('w:lineRule="exact"', image_paragraph._element.xml)


if __name__ == "__main__":
    unittest.main()
