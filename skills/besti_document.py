"""Deterministic Besti-style Chinese official DOCX generation."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BODY_SIZE = Pt(16)
TITLE_SIZE = Pt(22)
TABLE_BODY_SIZE = Pt(12)
TABLE_HEADER_SIZE = Pt(14)
EXACT_LINE_SPACING = Pt(29)
FIRST_LINE_INDENT = Pt(32)


def _set_run_font(
    run: Any,
    *,
    east_asia: str,
    size: Any,
    bold: bool = False,
    color: str | None = None,
) -> None:
    run.font.name = "Times New Roman"
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = _rgb(color)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), "Times New Roman")


def _rgb(value: str) -> Any:
    from docx.shared import RGBColor

    normalized = str(value or "000000").replace("#", "").upper()
    return RGBColor.from_string(normalized)


def _set_exact_line_spacing(paragraph: Any) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = EXACT_LINE_SPACING
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def _add_text_run(
    paragraph: Any,
    text: str,
    *,
    east_asia: str = "仿宋_GB2312",
    size: Any = BODY_SIZE,
    bold: bool = False,
    color: str | None = None,
) -> Any:
    run = paragraph.add_run(text)
    _set_run_font(
        run,
        east_asia=east_asia,
        size=size,
        bold=bold,
        color=color,
    )
    return run


def _set_cell_border(cell: Any, **edges: dict[str, str]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge_name, edge_data in edges.items():
        edge = tc_borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tc_borders.append(edge)
        edge.set(qn("w:val"), edge_data.get("val", "single"))
        edge.set(qn("w:sz"), edge_data.get("sz", "4"))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), edge_data.get("color", "000000"))


def _set_cell_margins(cell: Any, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _add_page_field(paragraph: Any) -> None:
    _add_text_run(paragraph, "— ", east_asia="宋体", size=Pt(14))
    run = paragraph.add_run()
    _set_run_font(run, east_asia="宋体", size=Pt(14))
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    _add_text_run(paragraph, " —", east_asia="宋体", size=Pt(14))


def _paragraph_role(text: str) -> tuple[str, str, bool]:
    stripped = text.strip()
    if re.match(r"^[一二三四五六七八九十百]+、", stripped):
        return "黑体", stripped, True
    if re.match(r"^（[一二三四五六七八九十百]+）", stripped):
        return "楷体_GB2312", stripped, False
    if re.match(r"^\d+[．.]", stripped):
        normalized = re.sub(r"^(\d+)[．.]\s*", r"\1．", stripped)
        return "仿宋_GB2312", normalized, False
    return "仿宋_GB2312", text, False


def _add_body_paragraph(document: Document, text: str) -> Any:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_exact_line_spacing(paragraph)
    font_name, normalized_text, bold = _paragraph_role(text)
    paragraph.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    _add_text_run(
        paragraph,
        normalized_text,
        east_asia=font_name,
        size=BODY_SIZE,
        bold=bold,
    )
    return paragraph


def _add_image(document: Document, root: Path, image: dict[str, Any]) -> None:
    raw_path = str(image.get("path") or "").strip()
    if not raw_path:
        return
    path = (root / raw_path).resolve()
    root_resolved = root.resolve()
    if root_resolved != path and root_resolved not in path.parents:
        raise RuntimeError(f"image path escapes workspace: {raw_path}")
    if not path.is_file():
        raise RuntimeError(f"image not found: {raw_path}")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    width_cm = max(3.0, min(float(image.get("width_cm") or 14.0), 15.6))
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(width_cm))

    caption = str(image.get("caption") or "").strip()
    if caption:
        paragraph.paragraph_format.keep_with_next = True
        caption_paragraph = document.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_exact_line_spacing(caption_paragraph)
        _add_text_run(
            caption_paragraph,
            caption,
            east_asia="仿宋_GB2312",
            size=Pt(14),
        )


def _add_attachment_marker(document: Document, index: int, total: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.page_break_before = True
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(29)
    label = "附件" if total == 1 else f"附件 {index}"
    _add_text_run(
        paragraph,
        label,
        east_asia="黑体",
        size=BODY_SIZE,
        bold=True,
    )


def _add_table(document: Document, table_spec: dict[str, Any]) -> None:
    title = str(table_spec.get("title") or "").strip()
    if title:
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_exact_line_spacing(title_paragraph)
        _add_text_run(
            title_paragraph,
            title,
            east_asia="方正小标宋简体",
            size=TITLE_SIZE,
        )
        spacer = document.add_paragraph()
        spacer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        spacer.paragraph_format.line_spacing = Pt(6)

    headers = [str(item) for item in table_spec.get("headers") or []]
    rows = [[str(value) for value in row] for row in table_spec.get("rows") or []]
    if not headers:
        return
    column_count = len(headers)
    normalized_rows = [(row + [""] * column_count)[:column_count] for row in rows]

    table = document.add_table(rows=1 + len(normalized_rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths = table_spec.get("column_widths_cm") or []
    if len(widths) != column_count:
        widths = [15.6 / column_count] * column_count

    all_rows = [headers, *normalized_rows]
    for row_index, (row, values) in enumerate(zip(table.rows, all_rows)):
        row.height = Cm(0.75)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for column_index, (cell, value) in enumerate(zip(row.cells, values)):
            cell.width = Cm(float(widths[column_index]))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = Pt(18)
            _add_text_run(
                paragraph,
                value,
                east_asia="黑体" if row_index == 0 else "仿宋_GB2312",
                size=TABLE_HEADER_SIZE if row_index == 0 else TABLE_BODY_SIZE,
                bold=row_index == 0,
            )

            top = "12" if row_index == 0 else "4"
            bottom = "12" if row_index == len(all_rows) - 1 else "4"
            left = "12" if column_index == 0 else "4"
            right = "12" if column_index == column_count - 1 else "4"
            _set_cell_border(
                cell,
                top={"sz": top},
                bottom={"sz": bottom},
                start={"sz": left},
                end={"sz": right},
                insideH={"sz": "4"},
                insideV={"sz": "4"},
            )


def generate_besti_docx(
    workspace_root: Path,
    *,
    output_path: Path,
    title: str,
    body: str,
    document_type: str = "white-paper",
    subtitle: str = "",
    recipient: str = "",
    organization: str = "",
    date_text: str = "",
    attachments: list[str] | None = None,
    images: list[dict[str, Any]] | None = None,
    tables: list[dict[str, Any]] | None = None,
    header_text: str = "",
    document_number: str = "",
) -> dict[str, Any]:
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.6)
    section.bottom_margin = Cm(3.0)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.7)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = BODY_SIZE
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")

    if document_type == "red-header":
        header_paragraph = document.add_paragraph()
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_exact_line_spacing(header_paragraph)
        _add_text_run(
            header_paragraph,
            header_text or organization,
            east_asia="方正小标宋简体",
            size=Pt(28),
            color="C00000",
        )
        if document_number:
            number_paragraph = document.add_paragraph()
            number_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_exact_line_spacing(number_paragraph)
            _add_text_run(
                number_paragraph,
                document_number,
                east_asia="仿宋_GB2312",
                size=BODY_SIZE,
            )

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_exact_line_spacing(title_paragraph)
    if document_type != "red-header":
        title_paragraph.paragraph_format.space_before = Pt(58)
    title_paragraph.paragraph_format.space_after = Pt(29)
    _add_text_run(
        title_paragraph,
        title,
        east_asia="方正小标宋简体",
        size=TITLE_SIZE,
    )

    if subtitle:
        subtitle_paragraph = document.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_exact_line_spacing(subtitle_paragraph)
        subtitle_paragraph.paragraph_format.space_after = Pt(29)
        _add_text_run(
            subtitle_paragraph,
            subtitle,
            east_asia="楷体_GB2312",
            size=BODY_SIZE,
        )

    if recipient:
        recipient_text = recipient.rstrip("：:") + "："
        recipient_paragraph = document.add_paragraph()
        _set_exact_line_spacing(recipient_paragraph)
        recipient_paragraph.paragraph_format.first_line_indent = Pt(0)
        _add_text_run(
            recipient_paragraph,
            recipient_text,
            east_asia="仿宋_GB2312",
            size=BODY_SIZE,
        )

    for line in str(body or "").replace("\r\n", "\n").split("\n"):
        if not line.strip():
            spacer = document.add_paragraph()
            _set_exact_line_spacing(spacer)
            continue
        _add_body_paragraph(document, line)

    table_specs = list(tables or [])
    body_tables = [
        table_spec
        for table_spec in table_specs
        if str(table_spec.get("placement") or "attachment").lower() == "body"
    ]
    attachment_tables = [
        table_spec
        for table_spec in table_specs
        if str(table_spec.get("placement") or "attachment").lower() != "body"
    ]
    for table_spec in body_tables:
        _add_table(document, table_spec)

    image_specs = list(images or [])
    body_images = [
        image
        for image in image_specs
        if str(image.get("placement") or "attachment").lower() == "body"
    ]
    attachment_images = [
        image
        for image in image_specs
        if str(image.get("placement") or "attachment").lower() != "body"
    ]
    for image in body_images:
        _add_image(document, workspace_root, image)

    attachment_items = [str(item).strip() for item in attachments or [] if str(item).strip()]
    for table_spec in attachment_tables:
        title = str(table_spec.get("title") or "").strip()
        if title and title not in attachment_items:
            attachment_items.append(title)
    for image in attachment_images:
        title = str(image.get("caption") or "").strip()
        if not title:
            title = Path(str(image.get("path") or "")).stem
        if title and title not in attachment_items:
            attachment_items.append(title)
    if attachment_items:
        spacer = document.add_paragraph()
        _set_exact_line_spacing(spacer)
        for index, item in enumerate(attachment_items, 1):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_exact_line_spacing(paragraph)
            paragraph.paragraph_format.left_indent = FIRST_LINE_INDENT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            prefix = "附件：" if len(attachment_items) == 1 else ("附件：" if index == 1 else "　　　")
            numbered = item if len(attachment_items) == 1 else f"{index}．{item}"
            _add_text_run(
                paragraph,
                prefix + numbered,
                east_asia="仿宋_GB2312",
                size=BODY_SIZE,
            )

    if organization or date_text:
        for _ in range(2):
            spacer = document.add_paragraph()
            _set_exact_line_spacing(spacer)
            spacer.paragraph_format.keep_with_next = True
        if organization:
            organization_paragraph = document.add_paragraph()
            organization_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_exact_line_spacing(organization_paragraph)
            organization_paragraph.paragraph_format.right_indent = Pt(64)
            organization_paragraph.paragraph_format.keep_with_next = True
            _add_text_run(
                organization_paragraph,
                organization,
                east_asia="仿宋_GB2312",
                size=BODY_SIZE,
            )
        if date_text:
            date_paragraph = document.add_paragraph()
            date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_exact_line_spacing(date_paragraph)
            date_paragraph.paragraph_format.right_indent = FIRST_LINE_INDENT
            _add_text_run(
                date_paragraph,
                date_text,
                east_asia="仿宋_GB2312",
                size=BODY_SIZE,
            )

    attachment_content_count = len(attachment_tables) + len(attachment_images)
    attachment_index = 0
    for table_spec in attachment_tables:
        attachment_index += 1
        _add_attachment_marker(document, attachment_index, attachment_content_count)
        _add_table(document, table_spec)

    for image in attachment_images:
        attachment_index += 1
        _add_attachment_marker(document, attachment_index, attachment_content_count)
        _add_image(document, workspace_root, image)

    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_field(footer_paragraph)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return {
        "path": output_path.relative_to(workspace_root.resolve()).as_posix(),
        "size": output_path.stat().st_size,
        "format": "Besti white-paper official document",
        "page_margins_cm": {"top": 3.6, "bottom": 3.0, "left": 2.7, "right": 2.7},
        "line_spacing_pt": 29,
        "images_inserted": len(image_specs),
        "tables_inserted": len(table_specs),
        "attachment_pages_started": attachment_content_count,
    }
